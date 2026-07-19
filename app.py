import io
import re
from collections import defaultdict
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import streamlit as st
import tifffile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from scipy.cluster.vq import kmeans2
from skimage import filters, measure, morphology

# ── Tunable detection constants ──────────────────────────────────────────────
BG_SIGMA = 40        # gaussian sigma for local background estimation
SIGNAL_THRESH = 200  # darkness signal threshold (raw 16-bit units)
MIN_BLOB_PX = 100    # blobs smaller than this are splash artifacts
N_ROWS, N_COLS = 4, 4

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="Colony Analyzer", page_icon="🔬", layout="wide")
st.title("🔬 Colony Plate Analyzer")
st.markdown(
    "Upload all plate images at once. Images are grouped by batch letter and "
    "sorted by date automatically."
)
st.warning(
    "**Upload the original 16-bit TIFF files straight from the gel imager.** "
    "PNG or JPG exports crush the dynamic range and destroy the measurement "
    "signal — colony values computed from them are not valid."
)

with st.sidebar:
    st.header("⚙️ Settings")
    st.markdown("**Filename format:** `A20260611.tif` → Batch A, June 11 2026")
    st.markdown("**Image format:** one cropped 4×4 grid per image, 16-bit TIFF")
    st.markdown("---")
    st.markdown(
        "**How to use:**\n"
        "1. Upload all images\n"
        "2. Click Analyze\n"
        "3. Download results per batch"
    )


# ── Filename parser ──────────────────────────────────────────────────────────
def parse_filename(filename):
    """Extract batch letter and date from e.g. A20260611.tif"""
    name = filename.rsplit(".", 1)[0]
    match = re.match(r"^([A-Za-z]+)(\d{8})$", name)
    if not match:
        return None, None
    batch = match.group(1).upper()
    try:
        date = datetime.strptime(match.group(2), "%Y%m%d")
    except ValueError:
        return None, None
    return batch, date


# ── Image loading (full bit depth preserved) ─────────────────────────────────
def load_gray(file_bytes):
    """
    Read an image at its native bit depth and return a float64 grayscale array.

    Critically this does NOT convert to uint8. The source TIFFs span roughly
    0-5000 in 16-bit units; casting to uint8 collapses them into ~5 brightness
    levels and destroys the colony signal entirely.
    """
    img = tifffile.imread(io.BytesIO(file_bytes))
    arr = np.asarray(img)
    if arr.ndim == 3:
        arr = arr.max(axis=2)
    return arr.astype(np.float64)


# ── Colony detection and measurement ─────────────────────────────────────────
def detect_colonies(gray):
    """
    Measure 16 colonies from a single cropped 4x4 grid image.

    1. Local background via large gaussian blur; signal = bg - gray, so a pixel
       is "dark" relative to its own neighbourhood. Cancels vignetting and
       uneven illumination, which a single global background value cannot.
    2. Threshold the signal image (not the raw image) to find blobs.
    3. Keep the 16 largest blobs. Later timepoints legitimately produce extra
       stray blobs, so this selection matters.
    4. Fit a 4x4 lattice by clustering row- and column-coordinates independently
       with k-means. Robust to stray blobs, unlike sort-and-chunk, which
       mis-assigns every colony after a single misplaced point.
    5. Integrate the darkness signal pixel by pixel within each colony mask.

    Returns a list of 16 dicts sorted R1C1..R4C4.
    """
    bg = filters.gaussian(gray, sigma=BG_SIGMA)
    signal = bg - gray

    smask = signal > SIGNAL_THRESH
    smask = morphology.remove_small_objects(smask, MIN_BLOB_PX)
    smask = morphology.closing(smask, morphology.disk(3))

    labels = measure.label(smask)
    props = measure.regionprops(labels, intensity_image=signal)
    props = sorted(props, key=lambda p: p.area, reverse=True)[: N_ROWS * N_COLS]

    if len(props) < 2:
        return [], signal, 0

    centroids = np.array([p.centroid for p in props], dtype=float)

    # Fit the lattice: cluster rows and columns independently.
    k_rows = min(N_ROWS, len(props))
    k_cols = min(N_COLS, len(props))
    row_centers, _ = kmeans2(centroids[:, 0], k_rows, minit="++", seed=1)
    col_centers, _ = kmeans2(centroids[:, 1], k_cols, minit="++", seed=1)
    row_centers = np.sort(row_centers)
    col_centers = np.sort(col_centers)

    colonies = []
    for prop, (cy, cx) in zip(props, centroids):
        r_idx = int(np.argmin(np.abs(row_centers - cy)))
        c_idx = int(np.argmin(np.abs(col_centers - cx)))
        mask = labels == prop.label
        colonies.append(
            {
                "colony_id": f"R{r_idx + 1}C{c_idx + 1}",
                "row": r_idx,
                "col": c_idx,
                "cy": float(cy),
                "cx": float(cx),
                "area_px": int(prop.area),
                "integrated_darkness": float(signal[mask].sum()),
                "mean_darkness": float(signal[mask].mean()),
            }
        )

    # Fill any lattice position that received no blob, so the table always has
    # a complete 4x4 and downstream code can rely on the colony_id set.
    found = {c["colony_id"] for c in colonies}
    for r in range(N_ROWS):
        for c in range(N_COLS):
            cid = f"R{r + 1}C{c + 1}"
            if cid not in found:
                colonies.append(
                    {
                        "colony_id": cid,
                        "row": r,
                        "col": c,
                        "cy": 0.0,
                        "cx": 0.0,
                        "area_px": 0,
                        "integrated_darkness": 0.0,
                        "mean_darkness": 0.0,
                    }
                )

    colonies.sort(key=lambda c: (c["row"], c["col"]))
    return colonies, signal, len(props)


# ── QC checks ────────────────────────────────────────────────────────────────
def run_qc(colony, all_colonies):
    flags = []
    if colony["area_px"] == 0:
        flags.append("colony not detected at this lattice position")
    dupes = [c for c in all_colonies if c["colony_id"] == colony["colony_id"]]
    if len(dupes) > 1:
        flags.append("duplicate lattice assignment")
    return ("FLAG" if flags else "OK"), "; ".join(flags)


# ── Chart builder ────────────────────────────────────────────────────────────
def make_chart(batch_data, metric, metric_label, batch_name):
    """Per-colony traces plus a bold mean line, against true elapsed days."""
    first_date = batch_data[0]["date"]
    days_since = [(d["date"] - first_date).days for d in batch_data]
    colony_ids = [c["colony_id"] for c in batch_data[0]["colonies"]]

    fig, ax = plt.subplots(figsize=(10, 6))

    series = []
    for cid in colony_ids:
        values = []
        for day in batch_data:
            colony = next(
                (c for c in day["colonies"] if c["colony_id"] == cid), None
            )
            values.append(
                float(colony[metric])
                if colony and colony["area_px"] > 0
                else np.nan
            )
        series.append(values)
        ax.plot(
            days_since, values, color="#4a9d4a", alpha=0.45, linewidth=1.2,
            marker="o", markersize=3,
        )

    mean_vals = np.nanmean(np.array(series, dtype=float), axis=0)
    ax.plot(
        days_since, mean_vals, color="black", linewidth=2.5, marker="o",
        markersize=5, label="Mean of 16 colonies", zorder=5,
    )

    ax.set_title(
        f"Batch {batch_name} — per-colony growth\n"
        f"{metric_label}, each thin line = one of 16 colonies",
        fontsize=12, fontweight="bold",
    )
    ax.set_xlabel("Days since first timepoint")
    ax.set_ylabel(f"{metric_label} (a.u.)")
    ax.set_xticks(sorted(set(days_since)))
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ── Word doc ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_word_doc(batch_name, batch_data, chart_bufs):
    doc = Document()
    title = doc.add_heading(f"Colony Analysis Report — Batch {batch_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dates_str = ", ".join(d["date"].strftime("%b %d %Y") for d in batch_data)
    doc.add_paragraph(f"Dates analyzed: {dates_str}")
    doc.add_paragraph("Colonies per plate: 16 (4×4 grid, technical replicates)")
    doc.add_paragraph(
        "Integrated darkness = sum of local-background-subtracted darkness "
        "over every pixel of a colony. Computed on 16-bit source data."
    )
    doc.add_paragraph("")

    for heading, buf in chart_bufs:
        doc.add_heading(heading, level=1)
        buf.seek(0)
        doc.add_picture(buf, width=Inches(6))
        doc.add_paragraph("")

    doc.add_heading("Raw Data Per Day", level=1)
    for day in batch_data:
        doc.add_heading(day["date"].strftime("%B %d, %Y"), level=2)

        valid = [c for c in day["colonies"] if c["area_px"] > 0]
        avg_id = np.mean([c["integrated_darkness"] for c in valid]) if valid else 0
        avg_area = np.mean([c["area_px"] for c in valid]) if valid else 0
        flags = sum(1 for c in day["colonies"] if c["qc_flag"] == "FLAG")
        doc.add_paragraph(
            f"Mean integrated darkness: {avg_id:,.0f}   |   "
            f"Mean area: {avg_area:,.0f} px   |   Flagged: {flags}/16"
        )

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(
            ["Colony", "Area (px)", "Integrated Darkness", "Mean Darkness",
             "QC", "Notes"]
        ):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(hdr[i], "D9E1F2")

        for c in day["colonies"]:
            rc = table.add_row().cells
            rc[0].text = c["colony_id"]
            rc[1].text = f"{c['area_px']:,}"
            rc[2].text = f"{c['integrated_darkness']:,.0f}"
            rc[3].text = f"{c['mean_darkness']:,.1f}"
            rc[4].text = c["qc_flag"]
            rc[5].text = c["qc_note"]
            set_cell_bg(rc[4], "FFD7D7" if c["qc_flag"] == "FLAG" else "D7FFD7")

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── CSV builder ──────────────────────────────────────────────────────────────
def build_csv(batch_name, batch_data):
    rows = []
    for day in batch_data:
        date_str = day["date"].strftime("%Y-%m-%d")
        for c in day["colonies"]:
            rows.append(
                {
                    "batch": batch_name,
                    "date": date_str,
                    "colony_id": c["colony_id"],
                    "area_px": c["area_px"],
                    "integrated_darkness": round(c["integrated_darkness"], 2),
                    "mean_darkness": round(c["mean_darkness"], 2),
                    "qc_flag": c["qc_flag"],
                    "qc_note": c["qc_note"],
                }
            )
        ids = [c["integrated_darkness"] for c in day["colonies"]]
        areas = [c["area_px"] for c in day["colonies"]]
        rows.append(
            {
                "batch": batch_name,
                "date": date_str,
                "colony_id": "AVERAGE",
                "area_px": round(float(np.mean(areas)), 1),
                "integrated_darkness": round(float(np.mean(ids)), 2),
                "mean_darkness": round(
                    float(np.mean([c["mean_darkness"] for c in day["colonies"]])), 2
                ),
                "std_integrated_darkness": round(float(np.std(ids)), 2),
                "std_area": round(float(np.std(areas)), 1),
            }
        )
    return pd.DataFrame(rows).to_csv(index=False).encode("utf-8")


# ── Main UI ──────────────────────────────────────────────────────────────────
st.markdown("### Upload Images")
st.info("Name files as `A20260611.tif` (letter + YYYYMMDD). Upload all batches together.")

uploaded_files = st.file_uploader(
    "Upload all plate images",
    type=["tif", "tiff"],
    accept_multiple_files=True,
)

run_btn = st.button(
    "🔬 Analyze All Batches", type="primary", use_container_width=True,
    disabled=not uploaded_files,
)

if run_btn and uploaded_files:
    batches = defaultdict(list)
    bad_files = []
    for f in uploaded_files:
        batch, date = parse_filename(f.name)
        if batch is None:
            bad_files.append(f.name)
        else:
            batches[batch].append((date, f))

    if bad_files:
        st.warning(
            f"Skipped files with unrecognized names: {', '.join(bad_files)}. "
            "Expected format: A20260611.tif"
        )

    if not batches:
        st.error("No valid files found. Check filenames match format: A20260611.tif")
        st.stop()

    for batch in batches:
        batches[batch] = sorted(batches[batch], key=lambda x: x[0])

    st.markdown(
        f"**Found {len(batches)} batch(es):** {', '.join(sorted(batches.keys()))}"
    )
    st.markdown("---")

    total_images = sum(len(v) for v in batches.values())
    progress = st.progress(0)
    processed = 0
    batch_results = []

    for batch_name in sorted(batches.keys()):
        batch_data = []
        status = st.empty()

        for date, file in batches[batch_name]:
            status.info(f"Batch {batch_name} — {date.strftime('%b %d %Y')}...")

            try:
                gray = load_gray(file.getvalue())
            except Exception as exc:
                st.error(f"Could not read {file.name}: {exc}")
                st.stop()

            colonies, _signal, n_blobs = detect_colonies(gray)
            if not colonies:
                st.error(
                    f"No colonies detected in {file.name}. Check the image is a "
                    "cropped 4×4 grid from the gel imager."
                )
                st.stop()

            for c in colonies:
                c["qc_flag"], c["qc_note"] = run_qc(c, colonies)

            batch_data.append(
                {"date": date, "colonies": colonies, "n_blobs": n_blobs}
            )
            processed += 1
            progress.progress(processed / total_images)

        status.empty()

        id_chart = make_chart(
            batch_data, "integrated_darkness", "Integrated darkness", batch_name
        )
        area_chart = make_chart(
            batch_data, "area_px", "Colony area (px)", batch_name
        )
        id_bytes = id_chart.getvalue()
        area_bytes = area_chart.getvalue()

        colony_ids = [c["colony_id"] for c in batch_data[0]["colonies"]]
        summary_rows = []
        for cid in colony_ids:
            ids, areas = [], []
            for day in batch_data:
                m = next(
                    (c for c in day["colonies"] if c["colony_id"] == cid), None
                )
                if m:
                    ids.append(m["integrated_darkness"])
                    areas.append(m["area_px"])
            summary_rows.append(
                {
                    "Colony": cid,
                    "Avg Integrated Darkness": round(np.mean(ids), 1) if ids else 0,
                    "Avg Area (px)": round(np.mean(areas), 1) if areas else 0,
                }
            )
        df_sum = pd.DataFrame(summary_rows)
        df_sum = pd.concat(
            [
                df_sum,
                pd.DataFrame(
                    [
                        {
                            "Colony": "AVERAGE",
                            "Avg Integrated Darkness": round(
                                df_sum["Avg Integrated Darkness"].mean(), 1
                            ),
                            "Avg Area (px)": round(df_sum["Avg Area (px)"].mean(), 1),
                        }
                    ]
                ),
            ],
            ignore_index=True,
        )

        csv_bytes = build_csv(batch_name, batch_data)
        word_bytes = build_word_doc(
            batch_name,
            batch_data,
            [
                ("Integrated Darkness Over Time", io.BytesIO(id_bytes)),
                ("Colony Area Over Time", io.BytesIO(area_bytes)),
            ],
        ).getvalue()

        batch_results.append(
            {
                "batch_name": batch_name,
                "id_chart_bytes": id_bytes,
                "area_chart_bytes": area_bytes,
                "df_summary": df_sum,
                "csv_bytes": csv_bytes,
                "word_bytes": word_bytes,
            }
        )

    st.session_state.batch_results = batch_results
    progress.progress(1.0)
    st.success("✅ All batches complete!")


# ── Display results (persists across reruns so downloads don't clear UI) ─────
if "batch_results" in st.session_state:
    for result in st.session_state.batch_results:
        batch_name = result["batch_name"]
        st.markdown(f"## Batch {batch_name}")

        col1, col2 = st.columns(2)
        with col1:
            st.image(
                result["id_chart_bytes"],
                caption=f"Batch {batch_name} — Integrated Darkness",
                use_column_width=True,
            )
        with col2:
            st.image(
                result["area_chart_bytes"],
                caption=f"Batch {batch_name} — Colony Area",
                use_column_width=True,
            )

        st.markdown(f"#### Colony Averages — Batch {batch_name}")
        st.dataframe(result["df_summary"], use_container_width=True, hide_index=True)

        st.markdown(f"#### Download — Batch {batch_name}")
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button(
                f"📥 CSV — Batch {batch_name}",
                result["csv_bytes"],
                file_name=f"batch_{batch_name}_analysis.csv",
                mime="text/csv",
                use_container_width=True,
                key=f"csv_{batch_name}",
            )
        with dl2:
            st.download_button(
                f"📄 Word Report — Batch {batch_name}",
                result["word_bytes"],
                file_name=f"batch_{batch_name}_report.docx",
                mime="application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
                use_container_width=True,
                key=f"word_{batch_name}",
            )

        st.markdown("---")
